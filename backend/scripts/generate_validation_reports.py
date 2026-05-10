"""Generate the 12-paper validation reports as Word .docx documents.

Step 5 of Sprint 4 (per SPRINT4_BOOTSTRAP_v2.md §7 + WorkingDoc §6).

For each of the 13 paper fixtures (12 papers + Leiva 2025 split into
Escombreras and Frövi sub-cases), this script:
  1. Builds the Case from the Q1-Q7 ground truth in WorkingDoc §3.1
  2. Runs the engine pipeline (l0 -> l1 -> pathway -> activate -> l2 -> l3)
  3. Writes a structured .docx with the 8-section canonical layout
     (Section 1 bibliographic; Section 2 case summary; Section 3
     methodological config adopted by authors; Section 4 Q1-Q7
     compilation; Section 5 tool-derived config; Section 6 paper-vs-
     tool comparison; Section 7 verdict; Section 8 full activation list)

Output:
  deliverables/T46/validation_reports/T46_VR_<short_ref>.docx

Usage (from repo root or from anywhere with PYTHONPATH set):
  cd backend
  $env:PYTHONPATH = "."
  python scripts/generate_validation_reports.py

The script is also covered by tests/test_validation_reports.py which
calls main() with a tmp_path output directory and verifies the files
exist and have the expected structure.
"""
from __future__ import annotations

import argparse
import sys
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from app.domain.enums import (
    Q1,
    Q2,
    Q4,
    Q5,
    Q7,
    Q6a,
    Q6b,
)
from app.domain.models import Q3, Case, Flow
from app.engine.pipeline import run as pipeline_run

# ---------------------------------------------------------------------------
# Per-paper data (Q1-Q7 + bibliographic metadata).
# Q1-Q7 mirrors test_12_papers_regression._PAPERS; metadata extracted
# from SYMBA_T46_Validation_WorkingDoc_v1.md §2.
# ---------------------------------------------------------------------------


def _flows(*qs: Q5) -> list[Flow]:
    return [Flow(id=f"f{i+1}", name=f"flow{i+1}", q5=q) for i, q in enumerate(qs)]


PAPERS: list[dict[str, Any]] = [
    {
        "id": "sokka_2011",
        "short_ref": "Sokka2011",
        "citation": "Sokka L. et al. (2011)",
        "country": "FI",
        "site": "Kymenlaakso forest industry complex (UPM Kymi paper mill + cogen + CaCO3 + ClO2 + WWTP + Kouvola town)",
        "primary_method": "LCA attributional, system expansion, black-box",
        "fu": "Annual production of UPM Kymi complex",
        "boundary": "Cradle-to-gate IES, six actors",
        "allocation": "System expansion (substitution credits)",
        "lcia": "ReCiPe midpoint, Finland-adjusted CFs",
        "scenarios": "3 reference scenarios (RS1 NG, RS2 peat, RS3 improvement)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True), "q4": {Q4.E},
        "flows": _flows(Q5.e),
        "q6a": Q6a.PULP_PAPER, "q6b": Q6b.TRL9, "q7": Q7.B,
        "rationale": (
            "Q1=B: eco-park (six-actor IES). Q2=D: existing baseline + RS3 "
            "improvement scenario. Q3=ENV: pure LCA. Q4=E: academic. Q5=e: "
            "black-box flow accounting. Q6a Other (Pulp & paper enum pending). "
            "Q6b TRL9, Q7=B regional."
        ),
    },
    {
        "id": "hashimoto_2010",
        "short_ref": "Hashimoto2010",
        "citation": "Hashimoto S. et al. (2010)",
        "country": "JP",
        "site": "Kawasaki Eco-town, D.C. Cement (Portland 348k t/y)",
        "primary_method": "LCCO2 (LCA reduced to CO2-only accounting)",
        "fu": "348,000 t Portland cement/year",
        "boundary": "Cradle-to-gate cement plant + raw mat. transport, calcination, combustion, waste transport+disposal",
        "allocation": "None (single-output)",
        "lcia": "Single-impact (CO2 only) using JP government EFs",
        "scenarios": "4 (S1 baseline no-IS, S2 current IS, S3 cluster-improved, S4 regional-extended)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True), "q4": {Q4.E},
        "flows": _flows(Q5.a, Q5.c),
        "q6a": Q6a.CEMENT_CONSTRUCTION, "q6b": Q6b.TRL9, "q7": Q7.B,
        "rationale": (
            "Q1=B: subject is the network's CO2 savings, not corporate ESG. "
            "Disambiguation rule 'who is the SUBJECT?' applies. Q2=D: "
            "S1+S2 baseline + S3+S4 alternatives. Q5 mixed (a slag pays-to-take, "
            "c by-products)."
        ),
    },
    {
        "id": "daddi_2017",
        "short_ref": "Daddi2017",
        "citation": "Daddi T. et al. (2017)",
        "country": "IT",
        "site": "Santa Croce sull'Arno tannery cluster (240 km², 6 municipalities)",
        "primary_method": "LCA attributional + system expansion (Mattila 2012)",
        "fu": "1 m2 finished bovine leather (cluster average), aligned with PCR EPD",
        "boundary": "Cradle-to-tannery-gate; farming, slaughterhouse, hide preservation, tanning, transport, WWT, waste",
        "allocation": "System expansion (Mattila 2012 approach)",
        "lcia": "ILCD recommended methods",
        "scenarios": "4 (S1 current IS, S2 hypothetical no-IS, S3a/S3b future water reuse)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.TEXTILE_LEATHER, "q6b": Q6b.TRL9, "q7": Q7.B,
        "rationale": "Cluster of 6 municipalities + WWT facility = Q1=B. S3a/S3b alternatives = Q2=D.",
    },
    {
        "id": "paulu_2022",
        "short_ref": "Paulu2022",
        "citation": "Paulu A. et al. (2022)",
        "country": "CZ",
        "site": "Czech Republic industry-wide CCP & C&DW management",
        "primary_method": "LCA sector-attributional, EF (PEF/OEF) normalised+weighted",
        "fu": "Annual national flows of CCP and C&DW",
        "boundary": "National-scale aggregated; transport from generation to use/disposal sites",
        "allocation": "Sector-level attributional, no per-flow allocation",
        "lcia": "Environmental Footprint (EF, normalised+weighted PEF/OEF)",
        "scenarios": "2 (current state + symbiotic optimised)",
        "q1": Q1.C, "q2": Q2.D, "q3": Q3(env=True), "q4": {Q4.D, Q4.E},
        "flows": _flows(Q5.e),
        "q6a": Q6a.WASTE_VALORIZATION, "q6b": Q6b.TRL9, "q7": Q7.D,
        "rationale": (
            "Pure policy assessment at national scale = Q1=C. Q4=D+E "
            "(EU PEF policy + academic publication, multi-select). "
            "Q7=D multi-scale industry-wide."
        ),
    },
    {
        "id": "arce_bastias_2023",
        "short_ref": "ArceBastias2023",
        "citation": "Arce Bastias F. et al. (2023)",
        "country": "AR",
        "site": "Madera Plástica Mendoza (3 entities, plastic recycling)",
        "primary_method": "LCA + new SPI indicator, three allocation methods tested",
        "fu": "Plastic waste managed by MPM network",
        "boundary": "Three-entity network, cradle-to-gate per entity",
        "allocation": "Three methods tested: full / partial (proposed SPI) / system expansion",
        "lcia": "CEENE (resources) + IPCC GWP100",
        "scenarios": "Symbiotic vs non-symbiotic per entity",
        "q1": Q1.B, "q2": Q2.A, "q3": Q3(env=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.PLASTICS_PACKAGING, "q6b": Q6b.TRL9, "q7": Q7.A,
        "rationale": (
            "3-entity network + ex-post analysis = Q1=B, Q2=A (operational). "
            "Methodological exploration paper (3 allocation methods) — tool "
            "applies system expansion default; expert override available."
        ),
    },
    {
        "id": "wiktor_2018",
        "short_ref": "Wiktor2018",
        "citation": "Wiktor & Johansson (2018)",
        "country": "SE",
        "site": "VA Syd Malmö sewage sludge (3 actors, 27,000 t/y)",
        "primary_method": "LCA + LCC integrated, Ecoinvent 3.3, marginal electricity = coal",
        "fu": "LCA: management of 27,000 t/y sludge × 1 year. LCC: 30 years discount 3.5%",
        "boundary": "Drying, incineration, P recovery (ASH DEC or EcoPhos), heat sources, capacity options",
        "allocation": "System expansion + entity-level cost allocation",
        "lcia": "Acidification + Eutrophication + GWP",
        "scenarios": "13 (S0 baseline + 12 combinations of 4 design choices)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True, eco=True), "q4": {Q4.E},
        "flows": _flows(Q5.a, Q5.c, Q5.c),
        "q6a": Q6a.WASTEWATER_BIOFACTORIES, "q6b": Q6b.TRL7_8, "q7": Q7.B,
        "rationale": (
            "Most extensive multi-scenario design (12 alternatives across 4 "
            "design dimensions) = Q2=D. Q6a wastewater_biofactories. Q6b 7-8 "
            "for new boiler design in -EC scenarios; per-scenario TRL specified "
            "in case data."
        ),
    },
    {
        "id": "leiva_2025_escombreras",
        "short_ref": "Leiva2025_Escombreras",
        "citation": "Leiva N. et al. (2025) - Escombreras demo",
        "country": "ES",
        "site": "Escombreras (Quimica del Estroncio KNO3 optimisation)",
        "primary_method": "MFCA (ISO 14051) + CBA + TEA, NO LCA",
        "fu": "1 hour KNO3 production",
        "boundary": "Symbiotic exchange links only (not full LCA cycle)",
        "allocation": "MFCA mass-based; CBA aggregates",
        "lcia": "OPEX, monetary efficiency, CAPEX, Sankey cost diagrams",
        "scenarios": "1+1 (baseline + post-IS)",
        "q1": Q1.A, "q2": Q2.D, "q3": Q3(eco=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.CHEMICALS_FERTILIZERS, "q6b": Q6b.TRL9, "q7": Q7.B,
        "rationale": (
            "Single-site process optimisation between QSr and neighbours = "
            "Q1=A. Pure ECO (MFCA + CBA + TEA, no LCA) = Q3 ECO-only. "
            "Validates the entire ECO-only branch."
        ),
    },
    {
        "id": "leiva_2025_frovi",
        "short_ref": "Leiva2025_Frovi",
        "citation": "Leiva N. et al. (2025) - Frovi demo",
        "country": "SE",
        "site": "Frovi (paper mill + 10 ha greenhouse + facilitator WA3RM)",
        "primary_method": "MFCA + CBA + TEA, multi-actor eco-park",
        "fu": "10 ha greenhouse / 100,000 m2",
        "boundary": "Symbiotic exchange links",
        "allocation": "MFCA mass-based",
        "lcia": "OPEX, CAPEX, Sankey diagrams",
        "scenarios": "4 (A CCU / B Scrubber / C External / D Baseline NG)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(eco=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.PULP_PAPER, "q6b": Q6b.TRL7_8, "q7": Q7.B,
        "rationale": "Multi-actor eco-park = Q1=B. CCU scenario at TRL 7-8.",
    },
    {
        "id": "danielsson_2018",
        "short_ref": "Danielsson2018",
        "citation": "Danielsson J. et al. (2018) - Symbiosis Center DK",
        "country": "DK",
        "site": "Kalundborg Symbiosis (8-member association since 1972)",
        "primary_method": "LCA consequential black-box + LCC analytical",
        "fu": "Year of production (Baseline 2015)",
        "boundary": "Black-box, three resource layers (energy / materials / water), 8 members",
        "allocation": "System expansion (substitution-based displacement)",
        "lcia": "t CO2, m3 water, MWh energy, t materials, M EUR savings",
        "scenarios": "3 (Baseline 2015, 2018, 2019)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True, eco=True), "q4": {Q4.C},
        "flows": _flows(Q5.e),
        "q6a": Q6a.ENERGY_UTILITIES, "q6b": Q6b.TRL9, "q7": Q7.B,
        "rationale": (
            "Most iconic IS network (8 actors, 50+ years operational) = Q1=B "
            "canonical. Q4=C public-claim (Symbiosis Center DK promotes "
            "Kalundborg as superior model)."
        ),
    },
    {
        "id": "kerdlap_2024",
        "short_ref": "Kerdlap2024",
        "citation": "Kerdlap P. et al. (2024)",
        "country": "(urban fictional)",
        "site": "Fictional urban agri-food network (5 entities)",
        "primary_method": "UM3-LCE3-ISN matrix-based LCA+LCC unified model",
        "fu": "Network output (illustrative)",
        "boundary": "Three-level (network / entity / specific resource flow)",
        "allocation": "Matrix-based decomposition",
        "lcia": "NPV (LCC) per entity + per flow + per network; LCA impacts at all 3 levels",
        "scenarios": "Open-loop and closed-loop variants",
        "q1": Q1.B, "q2": Q2.C, "q3": Q3(env=True, eco=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.AGRI_FOOD, "q6b": Q6b.TRL9, "q7": Q7.A,
        "rationale": (
            "Fictional 5-entity network = Q1=B. Methodological demonstration "
            "in design phase = Q2=C. NOTE: WorkingDoc S3.3 mapping table lists "
            "IS-01 but its own derivation rule 'Q2=C -> IS-04 (any Q1)' plus "
            "the gamma matrix (ADR-005) both yield IS-04. Engine follows the "
            "rule (IS-04)."
        ),
    },
    {
        "id": "subramanian_2021",
        "short_ref": "Subramanian2021",
        "citation": "Subramanian K. et al. (2021)",
        "country": "USA",
        "site": "The Plant, Chicago - multi-tenant building + on-site bread oven",
        "primary_method": "Capital-based LCSA (8 capitals, alternative to TBL)",
        "fu": "Bread baked in on-site oven",
        "boundary": "Multi-tenant building, on-site oven + AD system",
        "allocation": "Capital flow allocation across 8 capital types",
        "lcia": "Stock and flow of 8 capital types (natural / human / social / financial / manufactured / intellectual / cultural / political-institutional)",
        "scenarios": "3 (S1 natural gas, S2 biogas symbiotic, S3 renewable electricity)",
        "q1": Q1.B, "q2": Q2.D, "q3": Q3(env=True, eco=True, soc=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.MULTI_TENANT_URBAN_BUILDING, "q6b": Q6b.TRL9, "q7": Q7.A,
        "rationale": (
            "Multi-tenant urban building = Q1=B. Q3 full LCSA. Capital-based "
            "framework requires expert-mode override 'alternative S-LCA "
            "framework: capital-based' (default is UNEP/SETAC + CIRCPACK)."
        ),
    },
    {
        "id": "zhu_2013",
        "short_ref": "Zhu2013",
        "citation": "Zhu J. (2013) - TU Delft master thesis",
        "country": "NL",
        "site": "Generic 2-firm IS exchange",
        "primary_method": "LCA + simplified LCC, 'Symbiosis Assessment Diagram'",
        "fu": "UF1: 1 t finished product B. UF2: 1 t waste managed from A",
        "boundary": "Cradle-to-gate or gate-to-gate (varies by scenario)",
        "allocation": "Tested across scenarios (system expansion, partial, etc.)",
        "lcia": "Total cost (LCC) + LCA impacts",
        "scenarios": "19 scenarios varying FU, transport distance, waste quality, prices, allocation",
        "q1": Q1.A, "q2": Q2.C, "q3": Q3(env=True, eco=True), "q4": {Q4.E},
        "flows": _flows(Q5.c),
        "q6a": Q6a.MULTI_SECTOR, "q6b": Q6b.TRL9, "q7": Q7.A,
        "rationale": (
            "Generic 2-firm exchange in design phase = Q1=A, Q2=C. Most "
            "scenario-rich case in sample (19 scenarios). Engine maps to "
            "IS-04 via gamma matrix (Q1=A + Q2=C)."
        ),
    },
    {
        "id": "briassoulis_2023",
        "short_ref": "Briassoulis2023",
        "citation": "Briassoulis D. et al. (2023)",
        "country": "EU multi-country",
        "site": "Bio-based polymers framework + post-consumer IS recovery",
        "primary_method": "Full LCSA (LCA + S-LCA + LCC + TEA)",
        "fu": "1 t bio-polymer / post-consumer cycle",
        "boundary": "Bio-refinery + IS network, post-consumer recovery",
        "allocation": "Multi-method depending on flow",
        "lcia": "LCA: carbon, acidification, eutrophication. S-LCA: jobs, skills, wellbeing, H&S. LCC: CAPEX, OPEX, NPV. TEA: ROI, payback.",
        "scenarios": "Multi-scenario, framework propositional",
        "q1": Q1.C, "q2": Q2.C, "q3": Q3(env=True, eco=True, soc=True), "q4": {Q4.E},
        "flows": _flows(Q5.b, Q5.c),
        "q6a": Q6a.BIOBASED_POLYMERS, "q6b": Q6b.TRL9, "q7": None,
        "rationale": (
            "Industry-wide bio-polymer framework = Q1=C. Propositional "
            "framework in design = Q2=C. Q1=C overrides Q2 in gamma matrix "
            "-> IS-02. Full LCSA with TEA addition. Q7 not applicable (n/a "
            "in WorkingDoc)."
        ),
    },
]


# ---------------------------------------------------------------------------
# Document writer
# ---------------------------------------------------------------------------


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    run_k = p.add_run(f"{key}: ")
    run_k.bold = True
    p.add_run(value)


def _format_q4(q4: set[Q4]) -> str:
    return ", ".join(sorted(q.value for q in q4)) or "(none)"


def _format_q3(q3: Q3) -> str:
    parts = []
    if q3.env:
        parts.append("ENV")
    if q3.eco:
        parts.append("ECO")
    if q3.soc:
        parts.append("SOC")
    return " + ".join(parts) or "(none)"


def _format_flows(flows: list[Flow]) -> str:
    if not flows:
        return "(none)"
    return ", ".join(f"{f.id}={f.q5.value}" for f in flows)


def _build_case(paper: dict[str, Any]) -> Case:
    return Case(
        q1=paper["q1"],
        q2=paper["q2"],
        q3=paper["q3"],
        q4=paper["q4"],
        q6a=paper["q6a"],
        q6b=paper["q6b"],
        q7=paper["q7"],
        flows=paper["flows"],
    )


def _generate_one(paper: dict[str, Any], result: Case, out_dir: Path) -> Path:
    """Build the .docx for one paper and write to disk."""
    doc = Document()

    # Tighter font for code-like content
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # ----- Title -----
    title = doc.add_heading(f"Validation Report — {paper['citation']}", level=0)
    title.runs[0].font.size = Pt(20)
    doc.add_paragraph(
        f"SYMBA T4.6 IS Assessment Tool — auto-generated validation report "
        f"for paper short-ref {paper['short_ref']}."
    ).italic = True

    # ----- Section 1 — Bibliographic + IS context -----
    _add_heading(doc, "1. Bibliographic reference and IS context", level=1)
    _add_kv(doc, "Citation", paper["citation"])
    _add_kv(doc, "Country", paper["country"])
    _add_kv(doc, "Site / IS network", paper["site"])

    # ----- Section 2 — Case summary -----
    _add_heading(doc, "2. Case study summary", level=1)
    doc.add_paragraph(
        f"This paper applies the {paper['primary_method']} approach to the "
        f"case described above. The functional unit is {paper['fu']}; the "
        f"system boundary is {paper['boundary']}; allocation follows "
        f"{paper['allocation']}; LCIA / indicators: {paper['lcia']}; "
        f"scenarios analysed: {paper['scenarios']}."
    )

    # ----- Section 3 — Methodological config adopted by authors -----
    _add_heading(doc, "3. Methodological configuration adopted by the authors", level=1)
    _add_kv(doc, "3.1 Primary method", paper["primary_method"])
    _add_kv(doc, "3.2 Functional unit", paper["fu"])
    _add_kv(doc, "3.3 System boundary", paper["boundary"])
    _add_kv(doc, "3.4 Allocation", paper["allocation"])
    _add_kv(doc, "3.5 LCIA / indicators", paper["lcia"])
    _add_kv(doc, "3.6 Reference scenarios", paper["scenarios"])

    # ----- Section 4 — Q1-Q7 compilation -----
    _add_heading(doc, "4. Compilation of the seven user-facing tool questions", level=1)
    _add_kv(doc, "Q1 (IS scenario archetype)", paper["q1"].value)
    _add_kv(doc, "Q2 (temporal stance)", paper["q2"].value)
    _add_kv(doc, "Q3 (sustainability dimensions)", _format_q3(paper["q3"]))
    _add_kv(doc, "Q4 (use of results, multi-select)", _format_q4(paper["q4"]))
    _add_kv(doc, "Q5 per-flow", _format_flows(paper["flows"]))
    _add_kv(doc, "Q6a sector",
            paper["q6a"].value if paper["q6a"] else "(not set)")
    _add_kv(doc, "Q6b TRL band",
            paper["q6b"].value if paper["q6b"] else "(not set)")
    _add_kv(doc, "Q7 geographic spread",
            paper["q7"].value if paper["q7"] else "(n/a)")
    p = doc.add_paragraph()
    p.add_run("Compilation rationale: ").bold = True
    p.add_run(paper["rationale"])

    # ----- Section 5 — Tool-derived configuration -----
    _add_heading(doc, "5. Methodological configuration derived by the tool", level=1)
    _add_kv(doc, "5.1 Pathway",
            f"{result.pathway_id.value if result.pathway_id else '—'}"
            f"{' (extended)' if result.is_01_extended else ''}")
    _add_kv(doc, "5.2 ILCD situation",
            result.ilcd_situation.value if result.ilcd_situation else "—")
    _add_kv(doc, "5.3 LCC type",
            result.lcc_type.value if result.lcc_type else "—")
    _add_kv(doc, "5.4 S-LCA activation",
            result.slca_activation_state.value
            if result.slca_activation_state else "—")
    _add_kv(doc, "5.5 Activated nodes (count)", str(len(result.activated_nodes)))
    _add_kv(doc, "5.6 L1 BLOCK fired",
            ", ".join(result.blocked_by) if result.blocked_by else "(none)")
    _add_kv(doc, "5.7 L2 violations (count)", str(len(result.rule_violations)))
    _add_kv(doc, "5.8 L3 CDPs surfaced (count)", str(len(result.cdp_flags)))

    # Per-pillar config snapshot (only non-empty)
    _add_heading(doc, "5.9 Pillar configurations (engine-written keys)", level=2)
    pillar_pairs = [
        ("LCA", result.lca),
        ("LCC", result.lcc),
        ("S-LCA", result.slca),
        ("Report", result.report),
        ("Governance", result.governance),
        ("Methodological charter", result.methodological_charter),
        ("Review", result.review),
        ("System", result.system),
    ]
    any_written = False
    for name, pillar in pillar_pairs:
        if not pillar:
            continue
        any_written = True
        _add_heading(doc, name, level=3)
        for k, v in pillar.items():
            doc.add_paragraph(f"  {k} = {v}")
    if not any_written:
        doc.add_paragraph("(no pillar values written; case may be partially configured)")

    # ----- Section 6 — Comparison paper vs tool -----
    _add_heading(doc, "6. Comparison: paper configuration vs tool configuration", level=1)
    doc.add_paragraph(
        "Match assessment is qualitative for this auto-generated draft. "
        "Per WorkingDoc §1.2 the validation philosophy is methodological "
        "configuration coherence, not numerical replication. Detailed "
        "match/mismatch analysis is left to the human-reviewed final reports."
    )

    # ----- Section 7 — Verdict -----
    _add_heading(doc, "7. Validation verdict (preliminary)", level=1)
    if result.blocked_by:
        verdict = "BLOCKED at L1 — case as compiled triggers the listed L1 BLOCK(s); revise inputs."
    elif result.rule_violations:
        verdict = (
            f"PARTIAL MATCH — pipeline runs end-to-end but reports "
            f"{len(result.rule_violations)} L2 violation(s); inspect for "
            "expert overrides or assertion adjustments."
        )
    else:
        verdict = (
            "MATCH — pipeline runs end-to-end with zero L1/L2 violations. "
            "Tool configuration is consistent with the paper's methodological "
            "framing as compiled."
        )
    _add_kv(doc, "Preliminary verdict", verdict)

    # ----- Section 8 — Full activated nodes appendix -----
    _add_heading(doc, "8. Appendix: full node activation list", level=1)
    grouped: dict[str, list[str]] = {}
    for nid in result.activated_nodes:
        prefix = nid.split("_")[0]
        grouped.setdefault(prefix, []).append(nid)
    counts = Counter({k: len(v) for k, v in grouped.items()})
    summary_line = ", ".join(f"{k}={n}" for k, n in counts.most_common())
    _add_kv(doc, "Counts by prefix", summary_line)
    for prefix in sorted(grouped):
        _add_heading(doc, f"{prefix} ({len(grouped[prefix])})", level=2)
        doc.add_paragraph(", ".join(grouped[prefix]))

    out_path = out_dir / f"T46_VR_{paper['short_ref']}.docx"
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


DEFAULT_OUT_DIR = (
    Path(__file__).resolve().parent.parent.parent
    / "deliverables" / "T46" / "validation_reports"
)


def main(out_dir: Path | None = None) -> list[Path]:
    """Generate one .docx per paper. Returns the list of written paths."""
    out = out_dir or DEFAULT_OUT_DIR
    out.mkdir(parents=True, exist_ok=True)

    # Schemas are loaded once (lru_cached inside pipeline.run)
    written: list[Path] = []
    for paper in PAPERS:
        case = _build_case(paper)
        pipeline_run(case)
        path = _generate_one(paper, case, out)
        written.append(path)
        print(f"  wrote {path.name}")
    print(f"Generated {len(written)} validation reports in {out}")
    return written


def _cli() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=DEFAULT_OUT_DIR,
        help=f"Output directory (default: {DEFAULT_OUT_DIR})",
    )
    args = parser.parse_args()
    main(args.out_dir)
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
