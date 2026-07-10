"""Tests for app.services.dcf_docx."""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.domain.enums import Q1, Q2, Q4, Q5, Q7, Q6a, Q6b
from app.domain.models import Q3, Case, Flow
from app.engine.dcf_compose import compose_dcf
from app.engine.pipeline import run as pipeline_run
from app.services.dcf_docx import EU_FOOTER, generate_dcf_docx_bytes


@pytest.fixture(scope="module")
def dcf_schema() -> dict:
    p = (
        Path(__file__).resolve().parents[1] / "app" / "schemas" / "dcf_schema.json"
    )
    with open(p) as f:
        return json.load(f)


@pytest.fixture(scope="module")
def mandates_census() -> dict:
    p = (
        Path(__file__).resolve().parents[1]
        / "coordination" / "dcf_mandates_census.json"
    )
    with open(p) as f:
        return json.load(f)


def _flows(*qs):
    return [Flow(id=f"f{i+1}", name=f"flow{i+1}", q5=q) for i, q in enumerate(qs)]


@pytest.fixture()
def payload_wiktor(schemas, dcf_schema, mandates_census):
    case = Case(
        q1=Q1.B, q2=Q2.D, q3=Q3(env=True, eco=True), q4={Q4.E},
        q6a=Q6a.WASTEWATER_SLUDGE_BIOFACTORIES, q6b=Q6b.TRL7_8, q7=Q7.B,
        flows=_flows(Q5.a, Q5.c, Q5.c),
    )
    pipeline_run(case, schemas)
    return compose_dcf(case, dcf_schema, mandates_census)


def test_returns_bytes(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    assert isinstance(blob, bytes)
    assert len(blob) > 5000


def test_docx_parses_back(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    doc = Document(BytesIO(blob))
    # at least a few paragraphs
    assert len(doc.paragraphs) > 5


def test_docx_contains_title(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    doc = Document(BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Data Collection File" in text or "SYMBA T4.6" in text


def test_docx_contains_pathway(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    doc = Document(BytesIO(blob))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    assert "IS-01" in all_text


def test_docx_contains_eu_footer(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    doc = Document(BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "101135562" in text
    assert "www.symbaproject.eu" in text or "Horizon" in text


def test_docx_contains_mandate_rows(payload_wiktor):
    blob = generate_dcf_docx_bytes(payload_wiktor)
    doc = Document(BytesIO(blob))
    # The mandate table should have many rows
    mandate_table = None
    for table in doc.tables:
        head_row = table.rows[0]
        head_text = " ".join(c.text for c in head_row.cells)
        if "Node ID" in head_text and "Statement" in head_text:
            mandate_table = table
            break
    assert mandate_table is not None
    # 90 mandates expected (or near-90), plus header
    assert len(mandate_table.rows) > 50


def test_eu_footer_text_constants():
    assert "101135562" in EU_FOOTER
    assert "www.symbaproject.eu" in EU_FOOTER
