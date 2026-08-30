"""Round-trip: content drawn in the Network Builder reaches the exports.

Phase 3 of the Network Builder. Before this, the xlsx came out as an
empty grid and the docx described what the diagram *would* show; the
rows stored by /api/dcf/{case_id}/data now flow into both.
"""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from app.domain.dcf_data import DcfData, DcfRow
from app.domain.enums import Q1, Q2, Q4, Q5, Q7, Q6a, Q6b
from app.domain.models import Q3, Case, Flow
from app.engine.dcf_compose import compose_dcf
from app.engine.pipeline import run as pipeline_run
from app.services.dcf_docx import generate_dcf_docx_bytes
from app.services.dcf_excel import render_xlsx


@pytest.fixture(scope="module")
def dcf_schema() -> dict:
    p = Path(__file__).resolve().parents[1] / "app" / "schemas" / "dcf_schema.json"
    with open(p) as f:
        return json.load(f)


@pytest.fixture(scope="module")
def mandates_census() -> dict:
    p = (
        Path(__file__).resolve().parents[1]
        / "coordination" / "dcf_mandates_census.json"
    )
    with open(p) as f:
        return json.load(f)


@pytest.fixture()
def payload(schemas, dcf_schema, mandates_census):
    case = Case(
        q1=Q1.B, q2=Q2.D, q3=Q3(env=True, eco=True), q4={Q4.E},
        q6a=Q6a.WASTEWATER_SLUDGE_BIOFACTORIES, q6b=Q6b.TRL7_8, q7=Q7.B,
        flows=[Flow(id="f1", name="sludge", q5=Q5.a)],
    )
    pipeline_run(case, schemas)
    return compose_dcf(case, dcf_schema, mandates_census)


@pytest.fixture()
def drawn_network() -> DcfData:
    return DcfData(
        case_id="case-1",
        rows_by_section={
            "actors": [
                DcfRow(
                    row_id="a1",
                    values={"actor.name": "WWTP Alfa", "actor.role": "producer"},
                ),
                DcfRow(
                    row_id="a2",
                    values={"actor.name": "Biogas Beta", "actor.role": "consumer"},
                ),
            ],
            "flow_matrix": [
                DcfRow(
                    row_id="f1",
                    values={
                        "flow.name": "sludge",
                        "flow.origin_actor_id": "a1",
                        "flow.dest_actor_id": "a2",
                        "flow.type": "waste",
                        "flow.unit": "t/y",
                    },
                )
            ],
        },
        layout={},
    )


def _sheet_text(ws) -> str:
    return "\n".join(
        str(c.value) for row in ws.iter_rows() for c in row if c.value is not None
    )


# ---------------------------------------------------------------------------
# xlsx
# ---------------------------------------------------------------------------


def test_actors_tab_carries_the_stored_rows(payload, drawn_network):
    wb = load_workbook(BytesIO(render_xlsx(payload, drawn_network)))
    text = _sheet_text(wb["Actors"])
    assert "WWTP Alfa" in text
    assert "Biogas Beta" in text


def test_the_derived_id_column_is_filled_from_the_row_id(payload, drawn_network):
    wb = load_workbook(BytesIO(render_xlsx(payload, drawn_network)))
    ws = wb["Actors"]
    # header row 4, first data row 5; actor.id is column A
    assert ws.cell(row=5, column=1).value == "a1"
    assert ws.cell(row=6, column=1).value == "a2"


def test_blank_rows_remain_for_offline_collection(payload, drawn_network):
    wb = load_workbook(BytesIO(render_xlsx(payload, drawn_network)))
    ws = wb["Actors"]
    # two stored rows then blanks: row 7 is past the stored ones
    assert ws.cell(row=7, column=2).value is None


def test_network_tab_lists_the_edges(payload, drawn_network):
    wb = load_workbook(BytesIO(render_xlsx(payload, drawn_network)))
    text = _sheet_text(wb["Network Diagram"])
    assert "WWTP Alfa" in text
    assert "Biogas Beta" in text
    assert "sludge" in text


def test_export_without_stored_content_is_still_an_empty_grid(payload):
    wb = load_workbook(BytesIO(render_xlsx(payload)))
    ws = wb["Actors"]
    assert ws.cell(row=6, column=1).value is None
    assert "No network drawn yet" in _sheet_text(wb["Network Diagram"])


def test_unwired_flows_do_not_become_edges(payload, drawn_network):
    drawn_network.rows_by_section["flow_matrix"].append(
        DcfRow(row_id="f2", values={"flow.name": "biogas"})
    )
    wb = load_workbook(BytesIO(render_xlsx(payload, drawn_network)))
    # present in the Flow Matrix tab (it is a declared row) …
    assert "biogas" in _sheet_text(wb["Flow Matrix"])
    # … but not in the edge list, since it has no endpoints
    assert "biogas" not in _sheet_text(wb["Network Diagram"])


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------


def _docx_text(blob: bytes) -> str:
    doc = Document(BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_docx_reports_the_drawn_network(payload, drawn_network):
    text = _docx_text(generate_dcf_docx_bytes(payload, "wiktor", drawn_network))
    assert "WWTP Alfa" in text
    assert "Biogas Beta" in text
    assert "2 actor(s) and 1 wired flow(s)" in text


def test_docx_falls_back_to_the_descriptive_prose(payload):
    text = _docx_text(generate_dcf_docx_bytes(payload, "wiktor"))
    assert "rendered interactively" in text
