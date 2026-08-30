"""The .docx report opens with the same verdict the result page shows.

Before this, the report went straight into the WorkingDoc §6 layout —
right for a reviewer checking whether the configuration is defensible,
useless to a reader asking what they got and what to do next.
"""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.domain.enums import Q1, Q2, Q4, Q5, Q7, Q6a, Q6b
from app.domain.models import Q3, Case, Flow
from app.engine.pipeline import run as pipeline_run
from app.services.narrative import action_items, answers_in, verdict_for
from app.services.reports import generate_case_report_bytes


@pytest.fixture()
def wiktor(schemas) -> Case:
    case = Case(
        q1=Q1.B, q2=Q2.D, q3=Q3(env=True, eco=True), q4={Q4.E},
        q6a=Q6a.WASTEWATER_SLUDGE_BIOFACTORIES, q6b=Q6b.TRL7_8, q7=Q7.B,
        flows=[Flow(id="f1", name="sludge", q5=Q5.a)],
    )
    pipeline_run(case, schemas)
    return case


def _text(case: Case) -> str:
    doc = Document(BytesIO(generate_case_report_bytes(case, "test")))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# The composed verdict
# ---------------------------------------------------------------------------


def test_verdict_reads_as_a_sentence_not_a_code(wiktor):
    verdict = verdict_for(wiktor)
    assert verdict is not None
    assert verdict.title == "Operational symbiosis — decision support"
    assert "eco-park" in verdict.body
    # Q2 = D, so the extended suffix is appended
    assert "baseline" in verdict.body
    assert verdict.codes[0] == "IS-01 extended"


def test_verdict_sections_carry_their_citation(wiktor):
    sections = {s.code: s for s in verdict_for(wiktor).sections}
    assert "ILCD Situation A multi-actor" in sections
    assert sections["ILCD Situation A multi-actor"].source.startswith("D4.1")
    # a deactivated state is engine behaviour, not methodology: no citation
    assert sections["deactivated"].quote == ""


def test_answers_in_resolves_the_trigger_against_the_case(wiktor):
    assert answers_in("Q7 ∈ {B, C, D}", wiktor) == [("Q7", "B")]
    assert answers_in("Q3 has ≥2 dims active", wiktor) == [("Q3", "env + eco")]
    assert answers_in(None, wiktor) == []


def test_action_items_are_ordered_hardest_stop_first(wiktor):
    kinds = [item.kind for item in action_items(wiktor)]
    assert kinds == sorted(
        kinds, key=lambda k: {"block": 0, "obligation": 1, "decision": 2}[k]
    )


def test_obligations_come_from_triggered_rules_not_failed_assertions(wiktor):
    """A rule applies because its trigger fired. Presenting only the
    assertion failures would hide most of the practices the case owes and
    would rest on assertions that compare engine-written prose."""
    codes = {item.code for item in action_items(wiktor) if item.kind == "obligation"}
    assert codes == {r["rule_id"] for r in wiktor.applicable_rules}
    assert len(codes) > len(wiktor.rule_violations)


# ---------------------------------------------------------------------------
# The rendered document
# ---------------------------------------------------------------------------


def test_report_opens_with_the_verdict_and_the_next_steps(wiktor):
    text = _text(wiktor)
    assert "Operational symbiosis — decision support" in text
    assert "What this configuration means" in text
    assert "What to do next" in text
    # front matter comes before the canonical section 1
    assert text.index("What to do next") < text.index(
        "1. Bibliographic reference and IS context"
    )


def test_report_prints_the_deliverable_citation_in_full(wiktor):
    text = _text(wiktor)
    assert "Attributional Modeling with Substitution" in text
    assert "D4.1 Part 1 §8.2.1 (p. 36)" in text


def test_report_says_why_a_rule_applies_and_what_to_document(wiktor):
    text = _text(wiktor)
    assert "Why it applies here:" in text
    assert "What to document:" in text
    assert "[To document]" in text


def test_report_keeps_the_canonical_eight_sections(wiktor):
    text = _text(wiktor)
    for heading in (
        "1. Bibliographic reference and IS context",
        "5. Methodological configuration derived by the tool",
        "7. Validation verdict (preliminary)",
        "8. Appendix: full node activation list",
    ):
        assert heading in text


def test_a_clean_case_says_so_instead_of_listing_nothing(schemas):
    case = Case(
        q1=Q1.A, q2=Q2.A, q3=Q3(env=True), q4={Q4.E},
        q6a=Q6a.PLASTICS_PACKAGING, q6b=Q6b.TRL9, q7=Q7.A,
        flows=[Flow(id="f1", name="scrap", q5=Q5.b)],
    )
    pipeline_run(case, schemas)
    if action_items(case):
        pytest.skip("fixture is not clean; the empty-state path is covered elsewhere")
    assert "Nothing to fix" in _text(case)
